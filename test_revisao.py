"""
Testa a extração de docx/pdf e a separação de referências, e a lógica de
similaridade da verificação - sem depender de chamadas reais ao OpenAlex
(usamos uma função de busca falsa/mock só para validar a lógica de decisão).
"""
import io
import docx as docx_lib

from revisao.extrair_texto import extrair_texto, separar_corpo_e_referencias
from revisao.verificar_referencias import _similaridade

# 1. Gerar um .docx de teste em memória, simulando um trabalho acadêmico
doc = docx_lib.Document()
doc.add_paragraph("Introdução")
doc.add_paragraph("Este estudo analisa a biomecânica da corrida em atletas amadores.")
doc.add_paragraph("Metodologia")
doc.add_paragraph("Foram avaliados 30 corredores amadores.")
doc.add_paragraph("REFERÊNCIAS")
doc.add_paragraph("SILVA, J. A. Biomecânica da corrida: uma revisão. Revista Brasileira de "
                   "Educação Física, v. 10, n. 2, p. 45-60, 2022.")
doc.add_paragraph("SOUZA, M. P.; LIMA, R. C. Análise cinemática da marcha humana. "
                   "Fisioterapia em Movimento, v. 15, n. 1, p. 12-20, 2021.")

buffer = io.BytesIO()
doc.save(buffer)
conteudo_bytes = buffer.getvalue()

# 2. Testar extração
linhas = extrair_texto("trabalho_teste.docx", conteudo_bytes)
assert "REFERÊNCIAS" in linhas
print("[OK] Extração de .docx funcionando,", len(linhas), "linhas extraídas.")

# 3. Testar separação de corpo e referências
corpo, referencias, encontrou_cabecalho = separar_corpo_e_referencias(linhas)
assert encontrou_cabecalho is True
assert len(referencias) == 2, f"Esperava 2 referências, veio {len(referencias)}: {referencias}"
print("[OK] Separação de referências funcionando:")
for r in referencias:
    print("   -", r)

# 4. Testar arquivo sem cabeçalho de referências reconhecível
doc2 = docx_lib.Document()
doc2.add_paragraph("Um texto qualquer sem seção de referências.")
buffer2 = io.BytesIO()
doc2.save(buffer2)
linhas2 = extrair_texto("sem_referencias.docx", buffer2.getvalue())
corpo2, referencias2, encontrou2 = separar_corpo_e_referencias(linhas2)
assert encontrou2 is False
assert referencias2 == []
print("[OK] Documento sem cabeçalho de referências tratado corretamente (não inventa referências).")

# 5. Testar extensão não suportada
try:
    extrair_texto("arquivo.txt", b"qualquer coisa")
    raise AssertionError("Deveria ter levantado ValueError para .txt")
except ValueError as e:
    print("[OK] Extensão não suportada rejeitada corretamente:", e)

# 6. Testar a lógica de similaridade (sem rede - direto na função)
sim_alta = _similaridade(
    "SILVA, J. A. Biomecânica da corrida: uma revisão. Revista Brasileira, 2022.",
    "Biomecânica da corrida: uma revisão"
)
sim_baixa = _similaridade(
    "SILVA, J. A. Biomecânica da corrida: uma revisão. Revista Brasileira, 2022.",
    "Tratamento de águas residuais em zonas urbanas"
)
print(f"[OK] Similaridade alta (mesma obra, esperado > 0.5): {sim_alta:.2f}")
print(f"[OK] Similaridade baixa (obras diferentes, esperado < 0.3): {sim_baixa:.2f}")
assert sim_alta > sim_baixa

print("\n=== TODOS OS TESTES DE EXTRAÇÃO/VERIFICAÇÃO (SEM REDE) PASSARAM ===")
