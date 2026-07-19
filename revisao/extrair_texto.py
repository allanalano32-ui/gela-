"""
Extração de texto de documentos (.docx / .pdf) para a aba Revisão do GELA.

Regra de ouro mantida: este módulo só extrai texto literal do arquivo enviado.
Nenhum conteúdo é gerado, resumido ou "corrigido" aqui.
"""

import re
import io
import docx
import pdfplumber

# Cabeçalhos comuns de seção de referências em trabalhos acadêmicos em português.
# Aceita uma nota entre parênteses opcional depois (ex: "REFERÊNCIAS (ABNT NBR 6023:2018)"),
# que é comum quando o autor anota o padrão normativo usado.
PADROES_CABECALHO_REFERENCIAS = re.compile(
    r"^\s*(refer[êe]ncias(\s+bibliogr[áa]ficas)?|bibliografia)\s*(\([^)]*\))?\s*$",
    re.IGNORECASE,
)


def extrair_texto(nome_arquivo, conteudo_bytes):
    """
    Recebe o nome do arquivo (pra saber a extensão) e os bytes do conteúdo.
    Retorna o texto completo extraído, como uma lista de linhas não vazias.
    Levanta ValueError para extensões não suportadas.
    """
    extensao = nome_arquivo.lower().rsplit(".", 1)[-1] if "." in nome_arquivo else ""

    if extensao == "docx":
        return _extrair_docx(conteudo_bytes)
    elif extensao == "pdf":
        return _extrair_pdf(conteudo_bytes)
    else:
        raise ValueError(
            f"Extensão '.{extensao}' não suportada. Envie um arquivo .docx ou .pdf."
        )


def extrair_paragrafos_com_formatacao(nome_arquivo, conteudo_bytes):
    """
    Como extrair_texto(), mas preserva quais trechos de cada parágrafo estão em
    negrito ou itálico. Necessário para checar regras da ABNT que exigem
    destaque tipográfico (ex: título em negrito ou itálico).

    Só funciona de verdade para .docx (o formato preserva essa informação por
    trecho de texto). Para .pdf, retorna a formatação como "desconhecida" em
    todos os parágrafos - a checagem ABNT vai avisar isso explicitamente,
    em vez de fingir que verificou a formatação.

    Retorna uma lista de dicts: {"texto": str, "trechos_destacados": [str, ...],
    "formatacao_disponivel": bool}
    """
    extensao = nome_arquivo.lower().rsplit(".", 1)[-1] if "." in nome_arquivo else ""

    if extensao == "docx":
        documento = docx.Document(io.BytesIO(conteudo_bytes))
        resultado = []
        for p in documento.paragraphs:
            texto = p.text.strip()
            if not texto:
                continue
            trechos_destacados = [
                r.text for r in p.runs if (r.bold or r.italic) and r.text.strip()
            ]
            resultado.append({
                "texto": texto,
                "trechos_destacados": trechos_destacados,
                "formatacao_disponivel": True,
            })
        return resultado
    elif extensao == "pdf":
        linhas = _extrair_pdf(conteudo_bytes)
        return [
            {"texto": linha, "trechos_destacados": [], "formatacao_disponivel": False}
            for linha in linhas
        ]
    else:
        raise ValueError(
            f"Extensão '.{extensao}' não suportada. Envie um arquivo .docx ou .pdf."
        )


def _extrair_docx(conteudo_bytes):
    documento = docx.Document(io.BytesIO(conteudo_bytes))
    linhas = [p.text.strip() for p in documento.paragraphs if p.text.strip()]
    return linhas


def _extrair_pdf(conteudo_bytes):
    linhas = []
    with pdfplumber.open(io.BytesIO(conteudo_bytes)) as pdf:
        for pagina in pdf.pages:
            texto_pagina = pagina.extract_text() or ""
            for linha in texto_pagina.split("\n"):
                if linha.strip():
                    linhas.append(linha.strip())
    return linhas


def separar_corpo_e_referencias(linhas, linhas_sao_paragrafos=False):
    """
    Divide as linhas do documento em (corpo_do_texto, lista_de_referencias),
    procurando um cabeçalho tipo "REFERÊNCIAS" ou "REFERÊNCIAS BIBLIOGRÁFICAS".

    linhas_sao_paragrafos: quando True (típico de .docx, onde cada parágrafo do
    Word já corresponde a uma referência), cada linha após o cabeçalho é
    tratada como uma referência individual, SEM tentar rejuntar linhas quebradas.
    Quando False (típico de .pdf, onde uma referência longa pode ter sido
    quebrada em várias linhas pela extração), aplica a heurística de rejuntar
    linhas que não começam no padrão SOBRENOME.

    Se não encontrar um cabeçalho reconhecível, devolve a lista de referências
    vazia e avisa isso explicitamente - não tenta adivinhar onde começam as
    referências sem um marcador claro.
    """
    for i, linha in enumerate(linhas):
        if PADROES_CABECALHO_REFERENCIAS.match(linha):
            corpo = linhas[:i]
            bloco_referencias = linhas[i + 1:]
            if linhas_sao_paragrafos:
                referencias = [r for r in bloco_referencias if len(r) > 15]
            else:
                referencias = _dividir_em_referencias_individuais(bloco_referencias)
            return corpo, referencias, True  # True = cabeçalho encontrado

    # Nenhum cabeçalho de referências reconhecido
    return linhas, [], False


def _dividir_em_referencias_individuais(bloco_linhas):
    """
    Junta linhas quebradas de uma mesma referência (comum em PDFs, onde uma
    referência ABNT longa é quebrada em várias linhas) e separa referências
    distintas. Heurística: uma nova referência normalmente começa com um
    sobrenome em maiúsculas seguido de vírgula (padrão ABNT: SOBRENOME, Nome).
    """
    inicio_referencia = re.compile(r"^[A-ZÀ-Ú][A-ZÀ-Ú\.\s]+,")

    referencias = []
    atual = []
    for linha in bloco_linhas:
        if inicio_referencia.match(linha) and atual:
            referencias.append(" ".join(atual).strip())
            atual = [linha]
        else:
            atual.append(linha)
    if atual:
        referencias.append(" ".join(atual).strip())

    return [r for r in referencias if len(r) > 15]  # descarta lixo residual muito curto
